#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取分析报告中的内容
"""

import os
import json
from docx import Document
import PyPDF2

def extract_docx(file_path):
    """提取 DOCX 文件内容"""
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))

        return "\n".join(text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def extract_pdf(file_path):
    """提取 PDF 文件内容"""
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)
    except Exception as e:
        return f"Error reading {file_path}: {str(e)}"

def main():
    reports_dir = "/home/user/automate-system/分析报告"
    output_file = "/home/user/automate-system/extracted_reports.json"

    results = {}

    for filename in os.listdir(reports_dir):
        file_path = os.path.join(reports_dir, filename)

        if filename.endswith('.docx'):
            print(f"Processing: {filename}")
            content = extract_docx(file_path)
            results[filename] = {
                "type": "docx",
                "content": content,
                "length": len(content)
            }
        elif filename.endswith('.pdf'):
            print(f"Processing: {filename}")
            content = extract_pdf(file_path)
            results[filename] = {
                "type": "pdf",
                "content": content,
                "length": len(content)
            }

    # 保存结果
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Extracted {len(results)} reports")
    print(f"✓ Saved to: {output_file}")

    # 输出摘要
    for filename, data in results.items():
        print(f"\n{filename}: {data['length']} characters")

if __name__ == "__main__":
    main()
