#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""读取docx文件内容"""

import docx
import sys

def read_docx(file_path):
    """读取docx文件并输出所有文本内容"""
    try:
        doc = docx.Document(file_path)

        print("=" * 80)
        print(f"文件：{file_path}")
        print("=" * 80)
        print()

        # 读取所有段落
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:  # 只输出非空段落
                print(text)
                print()

        # 读取所有表格
        if doc.tables:
            print("\n" + "=" * 80)
            print("表格内容：")
            print("=" * 80)
            for table_idx, table in enumerate(doc.tables, 1):
                print(f"\n表格 {table_idx}:")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(" | ".join(row_data))
                print()

    except Exception as e:
        print(f"读取文件时出错: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python3 read_docx.py <docx文件路径>")
        sys.exit(1)

    read_docx(sys.argv[1])
